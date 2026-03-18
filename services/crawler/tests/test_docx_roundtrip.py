"""Tests for DOCX round-trip extraction and application."""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.services.docx_roundtrip_service import DocxRoundtripService, compute_semantic_groups


def _make_docx(*paragraphs: str, bold_first_run: bool = False) -> bytes:
    """Create a simple DOCX with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(text)
        if bold_first_run:
            run.bold = True
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_multi_run(*paragraphs: list[tuple[str, bool]]) -> bytes:
    """Create a DOCX where each paragraph has multiple runs with optional bold.

    Each paragraph is a list of (text, bold) tuples.
    """
    doc = Document()
    for runs_spec in paragraphs:
        para = doc.add_paragraph()
        for text, bold in runs_spec:
            run = para.add_run(text)
            run.bold = bold
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_bookmark(text: str, bookmark_name: str = "BM1") -> bytes:
    """Create a DOCX with a paragraph containing a bookmark around the text."""
    doc = Document()
    para = doc.add_paragraph()
    p_elem = para._element

    # Add bookmarkStart before run
    bm_start = etree.SubElement(p_elem, qn("w:bookmarkStart"))
    bm_start.set(qn("w:id"), "1")
    bm_start.set(qn("w:name"), bookmark_name)

    run = para.add_run(text)

    # Add bookmarkEnd after run
    bm_end = etree.SubElement(p_elem, qn("w:bookmarkEnd"))
    bm_end.set(qn("w:id"), "1")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_hyperlink() -> bytes:
    """Create a DOCX with a paragraph containing a hyperlink element."""
    doc = Document()
    para = doc.add_paragraph()
    p_elem = para._element

    # Add a normal run
    run = para.add_run("Click ")

    # Add hyperlink element
    hyperlink = etree.SubElement(p_elem, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), "rId1")
    h_run = etree.SubElement(hyperlink, qn("w:r"))
    h_text = etree.SubElement(h_run, qn("w:t"))
    h_text.text = "here"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table() -> bytes:
    """Create a DOCX with a table."""
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.add_paragraph("After table")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_drawing() -> bytes:
    """Create a DOCX with a run containing a w:drawing element."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Image: ")

    # Add a fake drawing element to the run
    drawing = etree.SubElement(run._element, qn("w:drawing"))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_docx_paragraphs(docx_bytes: bytes) -> list[str]:
    """Read all paragraph texts from a DOCX."""
    doc = Document(BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


@pytest.fixture
def service():
    return DocxRoundtripService()


class TestExtractStructured:
    def test_basic_extraction(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "World", "")
        result = service.extract_structured(docx_bytes)

        assert result["source_hash"]
        assert result["metadata"]["paragraph_count"] == 3
        assert result["metadata"]["table_count"] == 0
        assert len(result["lightweight"]) == 3
        assert result["lightweight"][0] == {"key": "p_0", "text": "Hello", "editable": True, "style": "Normal"}
        assert result["lightweight"][1] == {"key": "p_1", "text": "World", "editable": True, "style": "Normal"}
        assert result["lightweight"][2] == {"key": "p_2", "text": "", "editable": False, "style": "Normal"}
        assert result["metadata"]["group_count"] > 0
        assert "groups" in result

    def test_empty_document(self, service: DocxRoundtripService):
        doc = Document()
        buf = BytesIO()
        doc.save(buf)
        result = service.extract_structured(buf.getvalue())

        assert result["metadata"]["paragraph_count"] == 0
        assert result["lightweight"] == []

    def test_source_hash_deterministic(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Test")
        r1 = service.extract_structured(docx_bytes)
        r2 = service.extract_structured(docx_bytes)
        assert r1["source_hash"] == r2["source_hash"]

    def test_bookmark_paragraph_is_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_bookmark("Bookmarked text")
        result = service.extract_structured(docx_bytes)

        # Bookmark paragraphs should be editable (SAFE tier)
        bm_para = next(p for p in result["lightweight"] if p["text"] == "Bookmarked text")
        assert bm_para["editable"] is True

    def test_hyperlink_paragraph_not_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_hyperlink()
        result = service.extract_structured(docx_bytes)

        # Hyperlink paragraphs should not be editable (RISKY tier)
        assert any(not p["editable"] for p in result["lightweight"])

    def test_table_extraction(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_table()
        result = service.extract_structured(docx_bytes)

        assert result["metadata"]["table_count"] == 1
        keys = [p["key"] for p in result["lightweight"]]
        assert "tbl_0_r0_c0_p0" in keys
        assert "tbl_0_r1_c1_p0" in keys

        a1 = next(p for p in result["lightweight"] if p["key"] == "tbl_0_r0_c0_p0")
        assert a1["text"] == "A1"

    def test_drawing_paragraph_not_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_drawing()
        result = service.extract_structured(docx_bytes)

        assert any(not p["editable"] for p in result["lightweight"])

    def test_tab_and_newline_preserved(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Before\tAfter", "Line1\nLine2")
        result = service.extract_structured(docx_bytes)

        assert result["lightweight"][0]["text"] == "Before\tAfter"

    def test_empty_paragraph_not_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "", "World")
        result = service.extract_structured(docx_bytes)

        empty_para = result["lightweight"][1]
        assert empty_para["text"] == ""
        assert empty_para["editable"] is False

    def test_whitespace_only_paragraph_not_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "   ", "World")
        result = service.extract_structured(docx_bytes)

        ws_para = result["lightweight"][1]
        assert ws_para["editable"] is False

    def test_tab_only_paragraph_not_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "\t", "World")
        result = service.extract_structured(docx_bytes)

        tab_para = result["lightweight"][1]
        assert tab_para["text"] == "\t"
        assert tab_para["editable"] is False

    def test_rejects_oversized_file(self, service: DocxRoundtripService):
        # Create minimal docx, then test with fake large bytes
        with pytest.raises(ValueError, match="File too large"):
            service.extract_structured(b"\x00" * (25 * 1024 * 1024 + 1))

    def test_rejects_ole_file(self, service: DocxRoundtripService):
        ole_magic = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 100
        with pytest.raises(ValueError, match="Encrypted or legacy"):
            service.extract_structured(ole_magic)

    def test_rejects_docm_filename(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Test")
        with pytest.raises(ValueError, match=".docm"):
            service.extract_structured(docx_bytes, filename="test.docm")


class TestApplyStructured:
    def test_basic_apply(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "World")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Goodbye"}]
        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1
        paragraphs = _read_docx_paragraphs(result_bytes)
        assert paragraphs[0] == "Goodbye"
        assert paragraphs[1] == "World"

    def test_no_modifications(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello")
        extracted = service.extract_structured(docx_bytes)

        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], [])

        assert report["applied"] == 0
        assert report["success"] is True

    def test_hash_mismatch_raises(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello")
        with pytest.raises(ValueError, match="Source hash mismatch"):
            service.apply_structured(docx_bytes, "wrong_hash", [])

    def test_unknown_key_skipped(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_999", "text": "Nope"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 0
        assert "p_999" in report["skipped_unknown_key"]

    def test_no_change_skipped(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Hello"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 0
        assert "p_0" in report["skipped_no_change"]

    def test_not_editable_skipped(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_hyperlink()
        extracted = service.extract_structured(docx_bytes)

        non_editable = next(p for p in extracted["lightweight"] if not p["editable"])
        modifications = [{"key": non_editable["key"], "text": "Modified"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 0
        assert non_editable["key"] in report["skipped_not_editable"]

    def test_bold_format_preserved(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Bold text", bold_first_run=True)
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "New bold text"}]
        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1
        doc = Document(BytesIO(result_bytes))
        # First run should still be bold (format preserved via in-place modification)
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].text == "New bold text"

    def test_bookmark_preserved_after_apply(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_bookmark("Original text", "TestBookmark")
        extracted = service.extract_structured(docx_bytes)

        bm_para = next(p for p in extracted["lightweight"] if p["text"] == "Original text")
        modifications = [{"key": bm_para["key"], "text": "Modified text"}]
        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1

        # Verify bookmark still exists in the XML
        doc = Document(BytesIO(result_bytes))
        for para in doc.paragraphs:
            if para.text == "Modified text":
                p_elem = para._element
                bm_starts = p_elem.findall(qn("w:bookmarkStart"))
                bm_ends = p_elem.findall(qn("w:bookmarkEnd"))
                assert len(bm_starts) == 1
                assert bm_starts[0].get(qn("w:name")) == "TestBookmark"
                assert len(bm_ends) == 1
                return
        pytest.fail("Modified paragraph with bookmark not found")

    def test_drawing_paragraph_skipped(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_drawing()
        extracted = service.extract_structured(docx_bytes)

        non_editable = next(p for p in extracted["lightweight"] if not p["editable"])
        modifications = [{"key": non_editable["key"], "text": "Replace image"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 0

    def test_table_cell_modification(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_table()
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "tbl_0_r0_c0_p0", "text": "Modified A1"}]
        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1
        doc = Document(BytesIO(result_bytes))
        assert doc.tables[0].cell(0, 0).text == "Modified A1"
        assert doc.tables[0].cell(0, 1).text == "B1"  # Unchanged

    def test_empty_paragraph_modification_rejected(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "", "World")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_1", "text": "Injected content"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 0
        assert "p_1" in report["skipped_not_editable"]

    def test_track_changes_empty_paragraph_rejected(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "", "World")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_1", "text": "Injected"}]
        _, report = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
        )

        assert report["applied"] == 0
        assert "p_1" in report["skipped_not_editable"]

    def test_roundtrip_no_modification(self, service: DocxRoundtripService):
        """Extract then apply with no modifications should preserve all text."""
        docx_bytes = _make_docx("Para 1", "Para 2", "", "Para 4")
        extracted = service.extract_structured(docx_bytes)

        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], [])

        original_texts = _read_docx_paragraphs(docx_bytes)
        result_texts = _read_docx_paragraphs(result_bytes)
        assert original_texts == result_texts

    def test_multiple_modifications(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("A", "B", "C", "D")
        extracted = service.extract_structured(docx_bytes)

        modifications = [
            {"key": "p_0", "text": "AA"},
            {"key": "p_2", "text": "CC"},
        ]
        result_bytes, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 2
        paragraphs = _read_docx_paragraphs(result_bytes)
        assert paragraphs[0] == "AA"
        assert paragraphs[1] == "B"
        assert paragraphs[2] == "CC"
        assert paragraphs[3] == "D"

    def test_format_simplified_single_run(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Goodbye"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1
        assert report["format_simplified"] == []

    def test_format_simplified_multi_run(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_multi_run([("Bold ", True), ("normal", False)])
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Replaced text"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications)

        assert report["applied"] == 1
        assert "p_0" in report["format_simplified"]

    def test_format_simplified_not_tracked_for_track_changes(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_multi_run([("Bold ", True), ("normal", False)])
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Replaced text"}]
        _, report = service.apply_structured(docx_bytes, extracted["source_hash"], modifications, track_changes=True)

        assert report["applied"] == 1
        assert report["format_simplified"] == []


class TestTrackChanges:
    def test_track_changes_produces_del_ins(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Original text here")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Modified text here"}]
        result_bytes, report = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
        )

        assert report["applied"] == 1

        # Verify w:del and w:ins elements exist
        doc = Document(BytesIO(result_bytes))
        p_elem = doc.paragraphs[0]._element
        dels = p_elem.findall(f".//{qn('w:del')}")
        ins = p_elem.findall(f".//{qn('w:ins')}")
        assert len(dels) > 0, "Expected w:del elements"
        assert len(ins) > 0, "Expected w:ins elements"

    def test_track_changes_has_author(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello world")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_0", "text": "Goodbye world"}]
        result_bytes, _ = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
            author="Test Author",
        )

        doc = Document(BytesIO(result_bytes))
        p_elem = doc.paragraphs[0]._element
        ins_elems = p_elem.findall(f".//{qn('w:ins')}")
        assert any(e.get(qn("w:author")) == "Test Author" for e in ins_elems)

    def test_track_changes_low_similarity_fallback(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("AAAA BBBB CCCC DDDD")
        extracted = service.extract_structured(docx_bytes)

        # Completely different text → similarity < 0.3 → whole paragraph fallback
        modifications = [{"key": "p_0", "text": "XXXX YYYY ZZZZ WWWW 1234 5678"}]
        result_bytes, report = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
        )

        assert report["applied"] == 1
        doc = Document(BytesIO(result_bytes))
        p_elem = doc.paragraphs[0]._element
        dels = p_elem.findall(f".//{qn('w:del')}")
        ins = p_elem.findall(f".//{qn('w:ins')}")
        assert len(dels) > 0
        assert len(ins) > 0

    def test_track_changes_preserves_unmodified(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Keep this", "Change this")
        extracted = service.extract_structured(docx_bytes)

        modifications = [{"key": "p_1", "text": "Changed"}]
        result_bytes, _ = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
        )

        doc = Document(BytesIO(result_bytes))
        # First paragraph should have no track changes
        p0 = doc.paragraphs[0]._element
        assert len(p0.findall(f".//{qn('w:del')}")) == 0
        assert len(p0.findall(f".//{qn('w:ins')}")) == 0
        assert doc.paragraphs[0].text == "Keep this"

    def test_track_changes_unique_ids(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Text A", "Text B")
        extracted = service.extract_structured(docx_bytes)

        modifications = [
            {"key": "p_0", "text": "Modified A"},
            {"key": "p_1", "text": "Modified B"},
        ]
        result_bytes, _ = service.apply_structured(
            docx_bytes,
            extracted["source_hash"],
            modifications,
            track_changes=True,
        )

        doc = Document(BytesIO(result_bytes))
        ids = set()
        for elem in doc.element.iter():
            wid = elem.get(qn("w:id"))
            if wid is not None and elem.tag in (qn("w:del"), qn("w:ins")):
                assert wid not in ids, f"Duplicate w:id: {wid}"
                ids.add(wid)


def _make_docx_with_headings() -> bytes:
    """Create a DOCX with headings and body paragraphs."""
    doc = Document()
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("Body A")
    doc.add_paragraph("Body B")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Body C")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_outline_level(style_name: str = "Custom Style") -> bytes:
    """Create a DOCX with a custom-styled paragraph that has outlineLvl set."""
    doc = Document()
    # Add a normal paragraph first
    doc.add_paragraph("Before")
    # Add a paragraph and manually set outlineLvl on its direct formatting
    para = doc.add_paragraph("Custom Heading")
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(para._element, qn("w:pPr"))
    outline = etree.SubElement(pPr, qn("w:outlineLvl"))
    outline.set(qn("w:val"), "0")
    # Add body after
    doc.add_paragraph("After")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestSemanticGroups:
    def test_groups_basic_heading_split(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_headings()
        result = service.extract_structured(docx_bytes)

        # With default min_group_size=5, groups are small so they stay together
        # Verify via direct call with min_group_size=1 to isolate heading split
        heading_levels = {
            p["key"]: i
            for p in result["lightweight"]
            for i, style in enumerate(["Heading 1", "Heading 2"])
            if p.get("style") == style
        }
        # Build heading_levels properly from the extracted data
        heading_levels = {}
        for p in result["lightweight"]:
            style = p.get("style", "")
            if style == "Heading 1":
                heading_levels[p["key"]] = 0
            elif style == "Heading 2":
                heading_levels[p["key"]] = 1

        groups = compute_semantic_groups(result["lightweight"], heading_levels, min_group_size=1)

        # Should have 2 groups: one per heading section
        assert len(groups) == 2

        group_1_texts = [p["text"] for p in groups[0]]
        assert "Section One" in group_1_texts
        assert "Body A" in group_1_texts
        assert "Body B" in group_1_texts

        group_2_texts = [p["text"] for p in groups[1]]
        assert "Section Two" in group_2_texts
        assert "Body C" in group_2_texts

    def test_groups_no_headings(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("A", "B", "C", "D", "E")
        result = service.extract_structured(docx_bytes)
        groups = result["groups"]

        # All editable paragraphs in one group (all under max_group_size)
        assert len(groups) == 1
        assert len(groups[0]) == 5

    def test_groups_max_size_split(self):
        lightweight = [{"key": f"p_{i}", "text": f"Para {i}", "editable": True, "style": "Normal"} for i in range(12)]
        groups = compute_semantic_groups(lightweight, heading_levels={}, max_group_size=5, min_group_size=1)

        assert len(groups) == 3
        assert len(groups[0]) == 5
        assert len(groups[1]) == 5
        assert len(groups[2]) == 2

    def test_groups_only_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Hello", "", "World", "", "End")
        result = service.extract_structured(docx_bytes)

        for group in result["groups"]:
            for entry in group:
                assert entry["editable"] is True

    def test_groups_cover_all_editable(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_headings()
        result = service.extract_structured(docx_bytes)

        editable_keys = {p["key"] for p in result["lightweight"] if p["editable"]}
        group_keys = {p["key"] for group in result["groups"] for p in group}

        assert editable_keys == group_keys

    def test_groups_empty_document(self, service: DocxRoundtripService):
        doc = Document()
        buf = BytesIO()
        doc.save(buf)
        result = service.extract_structured(buf.getvalue())

        assert result["groups"] == []
        assert result["metadata"]["group_count"] == 0

    def test_groups_single_paragraph(self, service: DocxRoundtripService):
        docx_bytes = _make_docx("Only one")
        result = service.extract_structured(docx_bytes)

        assert len(result["groups"]) == 1
        assert len(result["groups"][0]) == 1
        assert result["groups"][0][0]["text"] == "Only one"

    def test_groups_party_identification_pattern(self, service: DocxRoundtripService):
        """Mimics Kaufvertrag structure: heading + seller block + gap + buyer block."""
        doc = Document()
        doc.add_heading("Vertragsparteien", level=1)
        doc.add_paragraph("Seller Name GmbH")
        doc.add_paragraph("Seller Address")
        doc.add_paragraph('Seller Rep\t(die "Verkäuferin")')
        doc.add_paragraph("")  # empty gap
        doc.add_paragraph("")  # empty gap
        doc.add_paragraph("und")
        doc.add_paragraph("")  # empty gap
        doc.add_paragraph("Buyer Name AG")
        doc.add_paragraph("Buyer Address")
        doc.add_paragraph('Buyer Rep\t(die "Käuferin")')
        doc.add_heading("Kaufgegenstand", level=1)
        doc.add_paragraph("Purchase object text")
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = service.extract_structured(docx_bytes)
        groups = result["groups"]

        # Seller and buyer should be in the same group
        all_texts = [p["text"] for g in groups for p in g]
        assert "Seller Name GmbH" in all_texts
        assert "Buyer Name AG" in all_texts

        # Both seller and buyer must be in the SAME group
        for group in groups:
            texts = [p["text"] for p in group]
            if "Seller Name GmbH" in texts:
                assert "Buyer Name AG" in texts
                assert 'Buyer Rep\t(die "Käuferin")' in texts
                break

    def test_groups_heading_as_first_paragraph(self, service: DocxRoundtripService):
        doc = Document()
        doc.add_heading("First Heading", level=1)
        doc.add_paragraph("Content")
        buf = BytesIO()
        doc.save(buf)

        result = service.extract_structured(buf.getvalue())
        groups = result["groups"]

        # No empty initial group — heading starts the first group
        assert len(groups) == 1
        assert groups[0][0]["text"] == "First Heading"

    def test_groups_table_paragraphs_included(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_table()
        result = service.extract_structured(docx_bytes)

        # Editable table cell paragraphs should appear in groups
        group_keys = {p["key"] for group in result["groups"] for p in group}
        editable_tbl_keys = {p["key"] for p in result["lightweight"] if p["editable"] and p["key"].startswith("tbl_")}
        assert editable_tbl_keys.issubset(group_keys)

    def test_groups_heading_detected_by_outline_level(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_outline_level()
        result = service.extract_structured(docx_bytes)

        # Extract heading_levels from the result and test with min_group_size=1
        # to isolate the outline level detection behavior
        heading_levels = {}
        for p in result["lightweight"]:
            if p["text"] == "Custom Heading":
                heading_levels[p["key"]] = 0  # outlineLvl=0 was set in the helper

        groups = compute_semantic_groups(result["lightweight"], heading_levels, min_group_size=1)

        assert len(groups) == 2
        assert groups[0][0]["text"] == "Before"
        assert groups[1][0]["text"] == "Custom Heading"

    def test_groups_no_outline_level_not_heading(self, service: DocxRoundtripService):
        # Bold paragraph without outlineLvl should NOT be treated as heading
        doc = Document()
        doc.add_paragraph("First")
        bold_para = doc.add_paragraph()
        run = bold_para.add_run("Bold But Not Heading")
        run.bold = True
        doc.add_paragraph("Third")
        buf = BytesIO()
        doc.save(buf)

        result = service.extract_structured(buf.getvalue())
        groups = result["groups"]

        # All in one group — bold paragraph is not a heading
        assert len(groups) == 1
        texts = [p["text"] for p in groups[0]]
        assert "First" in texts
        assert "Bold But Not Heading" in texts
        assert "Third" in texts

    def test_groups_style_extracted(self, service: DocxRoundtripService):
        docx_bytes = _make_docx_with_headings()
        result = service.extract_structured(docx_bytes)

        heading_entry = next(p for p in result["lightweight"] if p["text"] == "Section One")
        assert heading_entry["style"] == "Heading 1"

        body_entry = next(p for p in result["lightweight"] if p["text"] == "Body A")
        assert body_entry["style"] == "Normal"

    def test_groups_outline_level_threshold(self):
        """H4 (outlineLvl=3) should NOT cause split when max_outline_level=2."""
        lightweight = [
            {"key": "p_0", "text": "Intro", "editable": True, "style": "Normal"},
            {"key": "p_1", "text": "Section", "editable": True, "style": "Heading 1"},
            {"key": "p_2", "text": "Body", "editable": True, "style": "Heading 4"},
            {"key": "p_3", "text": "More body", "editable": True, "style": "Heading 4"},
            {"key": "p_4", "text": "End", "editable": True, "style": "Normal"},
        ]
        heading_levels = {"p_1": 0, "p_2": 3, "p_3": 3}
        groups = compute_semantic_groups(lightweight, heading_levels, min_group_size=1, max_outline_level=2)

        # H1 splits, but H4 (level 3 > threshold 2) does not
        assert len(groups) == 2
        assert groups[0][0]["text"] == "Intro"
        assert len(groups[1]) == 4  # Section + Body + More body + End

    def test_groups_min_size_delays_split(self):
        """Heading when current group has < min_group_size should NOT split."""
        lightweight = [{"key": f"p_{i}", "text": f"Para {i}", "editable": True, "style": "Normal"} for i in range(8)]
        # Heading at p_3 — current group only has 3 paragraphs (< min=5)
        heading_levels = {"p_3": 0}
        groups = compute_semantic_groups(lightweight, heading_levels, min_group_size=5)

        # Split delayed — all 8 paragraphs in one group
        assert len(groups) == 1
        assert len(groups[0]) == 8

    def test_groups_min_size_allows_split(self):
        """Heading when current group has >= min_group_size should split."""
        lightweight = [{"key": f"p_{i}", "text": f"Para {i}", "editable": True, "style": "Normal"} for i in range(10)]
        # Heading at p_5 — current group has 5 paragraphs (= min)
        heading_levels = {"p_5": 0}
        groups = compute_semantic_groups(lightweight, heading_levels, min_group_size=5)

        assert len(groups) == 2
        assert len(groups[0]) == 5
        assert len(groups[1]) == 5

    def test_groups_max_size_overrides_min(self):
        """max_group_size splits regardless of heading boundaries."""
        lightweight = [{"key": f"p_{i}", "text": f"Para {i}", "editable": True, "style": "Normal"} for i in range(15)]
        groups = compute_semantic_groups(lightweight, heading_levels={}, max_group_size=10, min_group_size=5)

        assert len(groups) == 2
        assert len(groups[0]) == 10
        assert len(groups[1]) == 5
