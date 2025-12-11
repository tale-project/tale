"""
DOCX Service for Word document generation.

Handles:
- DOCX generation from structured content
- DOCX generation from template + content (preserves headers, footers, styles)
"""

import logging
from copy import deepcopy
from io import BytesIO
from typing import Any, Dict, Optional

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocxService:
    """Service for generating DOCX documents."""

    def __init__(self):
        self._http_client: Optional[httpx.AsyncClient] = None

    async def _get_http_client(self) -> httpx.AsyncClient:
        """Get or create HTTP client."""
        if self._http_client is None:
            self._http_client = httpx.AsyncClient(timeout=60.0)
        return self._http_client

    async def cleanup(self):
        """Cleanup resources."""
        if self._http_client:
            await self._http_client.aclose()
            self._http_client = None

    async def download_file(self, url: str) -> bytes:
        """Download a file from URL."""
        client = await self._get_http_client()
        response = await client.get(url)
        response.raise_for_status()
        return response.content

    async def generate_docx(
        self,
        content: Dict[str, Any],
    ) -> bytes:
        """
        Generate a DOCX document from structured content.

        Args:
            content: Document content structure:
                {
                    "title": "Document Title",
                    "subtitle": "Optional subtitle",
                    "sections": [
                        {"type": "heading", "level": 1, "text": "Section Title"},
                        {"type": "paragraph", "text": "Paragraph text..."},
                        {"type": "bullets", "items": ["Item 1", "Item 2"]},
                        {"type": "numbered", "items": ["First", "Second"]},
                        {"type": "table", "headers": [...], "rows": [[...], [...]]},
                    ]
                }

        Returns:
            Generated DOCX as bytes
        """
        doc = Document()

        # Add title
        title_text = content.get("title", "Untitled Document")
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle if present
        subtitle_text = content.get("subtitle")
        if subtitle_text:
            subtitle = doc.add_paragraph(subtitle_text)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacer

        # Process sections
        sections = content.get("sections", [])
        for section in sections:
            self._process_section(doc, section)

        # Save to bytes
        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _try_apply_style(self, doc: Document, paragraph, style_name: str) -> None:
        """Apply a style if it exists in the document."""
        try:
            doc.styles[style_name]
            paragraph.style = style_name
        except KeyError:
            pass

    def _process_section(self, doc: Document, section: Dict[str, Any]) -> None:
        """Process a single section and add it to the document."""
        section_type = section.get("type", "paragraph")

        if section_type == "heading":
            level = section.get("level", 1)
            doc.add_heading(section.get("text", ""), level=level)

        elif section_type == "paragraph":
            doc.add_paragraph(section.get("text", ""))

        elif section_type == "bullets":
            for item in section.get("items", []):
                para = doc.add_paragraph(item)
                self._try_apply_style(doc, para, "List Bullet")

        elif section_type == "numbered":
            for item in section.get("items", []):
                para = doc.add_paragraph(item)
                self._try_apply_style(doc, para, "List Number")

        elif section_type == "table":
            self._add_table(doc, section)

        elif section_type == "quote":
            para = doc.add_paragraph(section.get("text", ""))
            self._try_apply_style(doc, para, "Quote")

        elif section_type == "code":
            doc.add_paragraph(section.get("text", ""))

    def _add_table(self, doc: Document, section: Dict[str, Any]) -> None:
        """Add a table to the document."""
        headers = section.get("headers", [])
        rows = section.get("rows", [])

        if not headers:
            return

        table = doc.add_table(rows=1, cols=len(headers))

        # Apply Table Grid style if available
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_data in rows:
            row = table.add_row()
            for i, cell_value in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = str(cell_value)

        doc.add_paragraph()  # Spacer after table

    # =========================================================================
    # TEMPLATE-BASED GENERATION
    # =========================================================================

    async def generate_docx_from_template(
        self,
        content: Dict[str, Any],
        template_bytes: bytes,
    ) -> bytes:
        """
        Generate a DOCX document using a template as the base.

        Preserves headers, footers, styles, and page setup from the template.
        """
        doc = Document(BytesIO(template_bytes))
        sections = content.get("sections", [])

        logger.info(
            f"Generating DOCX from template: title={content.get('title')}, "
            f"sections={len(sections)}"
        )

        self._clear_document_content(doc)

        # Add title
        if content.get("title"):
            title = doc.add_heading(content["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        if content.get("subtitle"):
            subtitle = doc.add_paragraph(content["subtitle"])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # Add sections
        for section in sections:
            self._process_section(doc, section)

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _clear_document_content(self, doc: Document) -> None:
        """
        Clear body content while preserving styles, headers, footers, and page setup.
        """
        first_section = doc.sections[0] if doc.sections else None

        # Preserve header/footer XML
        header_xml = None
        footer_xml = None
        if first_section:
            if first_section.header._element is not None:
                header_xml = deepcopy(first_section.header._element)
            if first_section.footer._element is not None:
                footer_xml = deepcopy(first_section.footer._element)

        # Remove paragraphs and tables from body
        for para in list(doc.paragraphs):
            para._element.getparent().remove(para._element)
        for table in list(doc.tables):
            table._tbl.getparent().remove(table._tbl)

        # Restore header/footer
        if doc.sections and first_section:
            section = doc.sections[0]
            self._restore_header_footer(section.footer, footer_xml)
            self._restore_header_footer(section.header, header_xml)

    def _restore_header_footer(self, header_footer, saved_xml) -> None:
        """Restore header or footer content from saved XML."""
        if saved_xml is None:
            return
        paras = saved_xml.findall('.//' + qn('w:p'))
        if paras:
            header_footer.is_linked_to_previous = False
            for p in list(header_footer._element):
                header_footer._element.remove(p)
            for p in paras:
                header_footer._element.append(deepcopy(p))


# Global service instance
_docx_service: Optional[DocxService] = None


def get_docx_service() -> DocxService:
    """Get or create the global DOCX service instance."""
    global _docx_service
    if _docx_service is None:
        _docx_service = DocxService()
    return _docx_service

