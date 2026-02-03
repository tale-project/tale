"""Smart DOCX text extraction with Vision API support.

This module extracts text and images from DOCX files while preserving
the relative positions of text and images in the document.
"""

import asyncio
from io import BytesIO

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from loguru import logger

from ...config import settings
from .openai_client import vision_client

MIN_IMAGE_SIZE = 10000  # ~100x100 pixels


async def _describe_image_bytes(
    image_bytes: bytes,
    semaphore: asyncio.Semaphore,
) -> str:
    """Describe image bytes using Vision API."""
    async with semaphore:
        try:
            if len(image_bytes) < MIN_IMAGE_SIZE:
                logger.debug(f"Skipping small image ({len(image_bytes)} bytes)")
                return ""
            return await vision_client.describe_image(image_bytes)
        except Exception as e:
            logger.warning(f"Failed to describe image: {e}")
            return ""


def _extract_table_text(table) -> str:
    """Extract text from a DOCX table."""
    rows_text = []
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        rows_text.append(" | ".join(cells_text))
    return "\n".join(rows_text)


async def extract_text_from_docx_bytes(
    docx_bytes: bytes,
    filename: str = "document.docx",
    *,
    process_images: bool = True,
) -> tuple[list[str], bool]:
    """Extract text from DOCX bytes with Vision support.

    Args:
        docx_bytes: Raw DOCX bytes
        filename: Optional filename for logging
        process_images: Whether to extract and describe embedded images

    Returns:
        Tuple of (list of content sections, vision_used flag)
    """
    logger.info(f"Processing DOCX: {filename}")

    doc = Document(BytesIO(docx_bytes))
    elements: list[tuple[int, str]] = []  # (position, content)
    vision_used = False
    position = 0
    semaphore = asyncio.Semaphore(settings.vision_max_concurrent_pages)

    # Collect image rels for later processing
    image_rels: dict[str, bytes] = {}
    if process_images:
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                try:
                    image_rels[rel.rId] = rel.target_part.blob
                except Exception as e:
                    logger.warning(f"Failed to extract image rel {rel.rId}: {e}")

    # Track which images have been processed (by rId)
    processed_image_rids: set[str] = set()

    # Process document body elements
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":  # paragraph
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break

            if para:
                text = para.text.strip()
                if text:
                    elements.append((position, text))
                    position += 1

                # Check for inline images in this paragraph
                if process_images:
                    # Look for blip elements (images) in the paragraph XML
                    blip_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
                    embed_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

                    for blip in element.iter(f"{blip_ns}blip"):
                        embed_id = blip.get(f"{embed_ns}embed")
                        if embed_id and embed_id in image_rels and embed_id not in processed_image_rids:
                            processed_image_rids.add(embed_id)
                            image_bytes = image_rels[embed_id]
                            description = await _describe_image_bytes(image_bytes, semaphore)
                            if description:
                                elements.append((position, f"[Image: {description}]"))
                                vision_used = True
                                position += 1

        elif tag == "tbl":  # table
            # Find the corresponding table object
            for table in doc.tables:
                if table._element == element:
                    table_text = _extract_table_text(table)
                    if table_text:
                        elements.append((position, f"[Table]\n{table_text}"))
                        position += 1
                    break

    # Process any remaining images that weren't inline
    if process_images:
        for rid, image_bytes in image_rels.items():
            if rid not in processed_image_rids:
                description = await _describe_image_bytes(image_bytes, semaphore)
                if description:
                    elements.append((position, f"[Image: {description}]"))
                    vision_used = True
                    position += 1

    # Sort by position (should already be in order, but just to be safe)
    elements.sort(key=lambda x: x[0])
    content_list = [elem[1] for elem in elements]

    logger.info(
        f"DOCX processing complete: {len(content_list)} elements, Vision API used: {vision_used}"
    )

    return content_list, vision_used
