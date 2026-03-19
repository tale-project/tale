"""
DOCX Round-Trip Service — extract structured JSON from DOCX and apply modifications back.

Uses python-docx for 100% structural fidelity. Modifies w:t elements in-place
to preserve bookmarks, comments, and all non-run XML children.
"""

from __future__ import annotations

import hashlib
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from loguru import logger

_MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB
_MAX_PARAGRAPHS = 5000
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# OOXML namespace
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Complex element classification
_RISKY_TAGS = frozenset({"hyperlink", "fldChar", "fldSimple", "ins", "del", "sdt"})
_CAUTION_TAGS = frozenset({"commentRangeStart", "commentRangeEnd"})
_SAFE_TAGS = frozenset({"bookmarkStart", "bookmarkEnd"})

# Non-text run content tags
_DRAWING_TAG = qn("w:drawing")
_BR_TAG = qn("w:br")

# Semantic grouping
_MAX_GROUP_SIZE = 10
_MIN_GROUP_SIZE = 5
_MAX_OUTLINE_LEVEL = 2  # Split at Heading 1-3 (outlineLvl 0-2), not H4+


def _tag_local(element) -> str:
    """Get the local tag name without namespace."""
    tag = element.tag
    if "}" in tag:
        return tag.split("}")[-1]
    return tag


def _compute_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def _check_file_safety(file_bytes: bytes, filename: str | None = None) -> None:
    """Reject encrypted, macro-enabled, or oversized files."""
    if len(file_bytes) > _MAX_FILE_SIZE:
        raise ValueError(f"File too large: {len(file_bytes)} bytes (max {_MAX_FILE_SIZE})")

    if file_bytes[:8] == _OLE_MAGIC:
        raise ValueError("Encrypted or legacy .doc file detected (OLE Compound Document)")

    if filename and filename.lower().endswith(".docm"):
        raise ValueError("Macro-enabled .docm files are not supported")


def _has_non_text_content(run_element) -> bool:
    """Check if a run contains drawings or page breaks."""
    for child in run_element:
        if child.tag == _DRAWING_TAG:
            return True
        if child.tag == _BR_TAG and child.get(qn("w:type")) == "page":
            return True
    return False


def _classify_paragraph(para_element) -> tuple[bool, list[str]]:
    """Classify a paragraph's complex elements.

    Returns (editable, detected_types).
    """
    detected: list[str] = []
    has_risky = False

    for child in para_element:
        local = _tag_local(child)
        if local in _RISKY_TAGS:
            detected.append(local)
            has_risky = True
        elif local in _CAUTION_TAGS or local in _SAFE_TAGS:
            detected.append(local)

    # Also check inside runs for non-text content
    for run_el in para_element.findall(qn("w:r")):
        if _has_non_text_content(run_el):
            has_risky = True
            detected.append("non_text_content")

    return (not has_risky, detected)


def _get_element_text(p_element) -> str:
    """Get full text from a w:p XML element by concatenating all w:t texts."""
    texts = []
    for t_el in p_element.iter(qn("w:t")):
        if t_el.text:
            texts.append(t_el.text)
    return "".join(texts)


def _iter_body_elements(doc: Document):
    """Iterate body elements, yielding (tag, element) for w:p and w:tbl only.

    Also handles w:sdt by extracting paragraphs from sdtContent.
    """
    for element in doc.element.body:
        tag = _tag_local(element)
        if tag == "p":
            yield ("p", element)
        elif tag == "tbl":
            yield ("tbl", element)
        elif tag == "sdt":
            # SDT can contain paragraphs in sdtContent
            sdt_content = element.find(qn("w:sdtContent"))
            if sdt_content is not None:
                for child in sdt_content:
                    child_tag = _tag_local(child)
                    if child_tag == "p":
                        yield ("sdt_p", child)
                    elif child_tag == "tbl":
                        yield ("sdt_tbl", child)


def _get_outline_level(para_element, para=None) -> int | None:
    """Get the outline level of a paragraph (0=Heading 1, 1=Heading 2, etc.).

    The w:outlineLvl element is the authoritative OOXML signal for heading
    status. Checks both direct paragraph formatting and the style definition.
    Returns None if no outline level is set.
    """
    # Check direct paragraph formatting
    pPr = para_element.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            return int(ol.get(qn("w:val")))
    # Check style definition
    if para is not None and para.style and para.style.element is not None:
        style_pPr = para.style.element.find(qn("w:pPr"))
        if style_pPr is not None:
            ol = style_pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                return int(ol.get(qn("w:val")))
    return None


def compute_semantic_groups(
    lightweight: list[dict[str, Any]],
    heading_levels: dict[str, int],
    max_group_size: int = _MAX_GROUP_SIZE,
    min_group_size: int = _MIN_GROUP_SIZE,
    max_outline_level: int = _MAX_OUTLINE_LEVEL,
) -> list[list[dict[str, Any]]]:
    """Group editable paragraphs into semantic batches.

    Single-pass algorithm:
    - Split at heading paragraphs (level <= max_outline_level) only when
      the current group has accumulated at least min_group_size paragraphs
    - Split when group reaches max_group_size regardless
    - Only include editable paragraphs in groups
    """
    groups: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []

    for entry in lightweight:
        if not entry.get("editable"):
            continue

        level = heading_levels.get(entry["key"])
        is_split_heading = level is not None and level <= max_outline_level

        if is_split_heading and len(current) >= min_group_size:
            groups.append(current)
            current = []

        current.append(entry)

        if len(current) >= max_group_size:
            groups.append(current)
            current = []

    if current:
        groups.append(current)

    return groups


class DocxRoundtripService:
    """Extract structured data from DOCX and apply text modifications back."""

    def extract_structured(self, file_bytes: bytes, filename: str | None = None) -> dict[str, Any]:
        """Extract structured paragraph data from a DOCX file.

        Returns:
            {
                "source_hash": "sha256hex",
                "metadata": {"paragraph_count": N, "table_count": N, "group_count": N},
                "lightweight": [{"key": "p_0", "text": "...", "editable": True, "style": "Normal"}, ...],
                "groups": [[...], ...]
            }
        """
        _check_file_safety(file_bytes, filename)

        source_hash = _compute_hash(file_bytes)
        doc = Document(BytesIO(file_bytes))

        # Build paragraph map for text access
        para_map = {p._element: p for p in doc.paragraphs}

        lightweight: list[dict[str, Any]] = []
        heading_levels: dict[str, int] = {}
        p_counter = 0
        tbl_counter = 0

        for tag, element in _iter_body_elements(doc):
            if tag in ("p", "sdt_p"):
                key = f"p_{p_counter}"
                p_counter += 1

                para = para_map.get(element)
                text = para.text if para else _get_element_text(element)
                style = para.style.name if para and para.style else None

                level = _get_outline_level(element, para)
                if level is not None:
                    heading_levels[key] = level

                if tag == "sdt_p" or not text.strip():
                    lightweight.append({"key": key, "text": text, "editable": False, "style": style})
                else:
                    editable, _detected = _classify_paragraph(element)
                    lightweight.append({"key": key, "text": text, "editable": editable, "style": style})

            elif tag in ("tbl", "sdt_tbl"):
                tbl_key = f"tbl_{tbl_counter}"
                tbl_counter += 1
                # Extract table cell paragraphs
                for row_idx, tr in enumerate(element.findall(qn("w:tr"))):
                    for col_idx, tc in enumerate(tr.findall(qn("w:tc"))):
                        for p_idx, p_el in enumerate(tc.findall(qn("w:p"))):
                            cell_key = f"{tbl_key}_r{row_idx}_c{col_idx}_p{p_idx}"
                            p_counter += 1
                            text = _get_element_text(p_el)
                            cell_para = para_map.get(p_el)
                            cell_style = cell_para.style.name if cell_para and cell_para.style else None
                            if not text.strip():
                                lightweight.append(
                                    {"key": cell_key, "text": text, "editable": False, "style": cell_style}
                                )
                            else:
                                editable, _detected = _classify_paragraph(p_el)
                                lightweight.append(
                                    {"key": cell_key, "text": text, "editable": editable, "style": cell_style}
                                )

        if p_counter > _MAX_PARAGRAPHS:
            raise ValueError(
                f"Document has {p_counter} paragraphs (max {_MAX_PARAGRAPHS}). Consider splitting the document."
            )

        groups = compute_semantic_groups(lightweight, heading_levels)

        logger.info(
            f"Extracted {p_counter} paragraphs, {tbl_counter} tables, "
            f"{len(groups)} groups from DOCX (hash={source_hash[:12]}...)"
        )

        return {
            "source_hash": source_hash,
            "metadata": {
                "paragraph_count": p_counter,
                "table_count": tbl_counter,
                "group_count": len(groups),
            },
            "lightweight": lightweight,
            "groups": groups,
        }

    def apply_structured(
        self,
        template_bytes: bytes,
        source_hash: str,
        modifications: list[dict[str, str]],
        track_changes: bool = False,
        author: str = "AI Assistant",
    ) -> tuple[bytes, dict[str, Any]]:
        """Apply text modifications to a DOCX template.

        Args:
            template_bytes: Original DOCX file bytes.
            source_hash: SHA-256 hash from extract step (for validation).
            modifications: List of {"key": "p_N", "text": "new text"}.
            track_changes: Whether to use Track Changes markup.
            author: Author name for Track Changes.

        Returns:
            (modified_docx_bytes, report)
        """
        actual_hash = _compute_hash(template_bytes)
        if actual_hash != source_hash:
            raise ValueError(
                f"Source hash mismatch: expected {source_hash[:16]}..., "
                f"got {actual_hash[:16]}.... The template file may have been modified."
            )

        doc = Document(BytesIO(template_bytes))
        para_map = {p._element: p for p in doc.paragraphs}

        # Build modification lookup
        mod_lookup: dict[str, str] = {}
        for mod in modifications:
            key = mod.get("key", "")
            text = mod.get("text", "")
            if key and isinstance(key, str) and isinstance(text, str):
                mod_lookup[key] = text

        report: dict[str, Any] = {
            "total_modifications_requested": len(modifications),
            "applied": 0,
            "skipped_not_editable": [],
            "skipped_unknown_key": [],
            "skipped_no_change": [],
            "skipped_non_text_content": [],
            "format_simplified": [],
            "errors": [],
        }

        # Track which keys were seen
        seen_keys: set[str] = set()

        # Track Changes writer (lazy init)
        tc_writer = None
        if track_changes:
            from app.services.docx_track_changes import TrackChangesWriter

            tc_writer = TrackChangesWriter(doc.element, author=author)

        p_counter = 0
        tbl_counter = 0

        for tag, element in _iter_body_elements(doc):
            if tag in ("p", "sdt_p"):
                key = f"p_{p_counter}"
                p_counter += 1

                if key not in mod_lookup:
                    continue

                seen_keys.add(key)
                new_text = mod_lookup[key]

                # SDT paragraphs are not editable
                if tag == "sdt_p":
                    report["skipped_not_editable"].append(key)
                    continue

                para = para_map.get(element) or Paragraph(element, None)

                old_text = para.text
                if not old_text.strip():
                    report["skipped_not_editable"].append(key)
                    continue

                editable, _detected = _classify_paragraph(element)
                if not editable:
                    report["skipped_not_editable"].append(key)
                    continue

                if old_text == new_text:
                    report["skipped_no_change"].append(key)
                    continue

                if not para.runs:
                    report["errors"].append({"key": key, "error": "paragraph has no runs"})
                    continue

                # Check for non-text content in runs
                has_non_text = False
                for run in para.runs:
                    if _has_non_text_content(run._element):
                        has_non_text = True
                        break

                if has_non_text:
                    report["skipped_non_text_content"].append(key)
                    continue

                try:
                    if tc_writer:
                        tc_writer.apply_paragraph_change(para, old_text, new_text)
                    else:
                        if len(para.runs) > 1:
                            report["format_simplified"].append(key)
                        _replace_paragraph_text(para, new_text)
                    report["applied"] += 1
                except Exception as e:
                    logger.warning(f"Failed to apply modification to {key}: {e}")
                    report["errors"].append({"key": key, "error": str(e)})

            elif tag in ("tbl", "sdt_tbl"):
                tbl_key = f"tbl_{tbl_counter}"
                tbl_counter += 1
                for row_idx, tr in enumerate(element.findall(qn("w:tr"))):
                    for col_idx, tc in enumerate(tr.findall(qn("w:tc"))):
                        for p_idx, p_el in enumerate(tc.findall(qn("w:p"))):
                            cell_key = f"{tbl_key}_r{row_idx}_c{col_idx}_p{p_idx}"
                            p_counter += 1

                            if cell_key not in mod_lookup:
                                continue

                            seen_keys.add(cell_key)
                            new_text = mod_lookup[cell_key]

                            para = para_map.get(p_el) or Paragraph(p_el, None)

                            old_text = para.text
                            if not old_text.strip():
                                report["skipped_not_editable"].append(cell_key)
                                continue

                            editable, _detected = _classify_paragraph(p_el)
                            if not editable:
                                report["skipped_not_editable"].append(cell_key)
                                continue

                            if old_text == new_text:
                                report["skipped_no_change"].append(cell_key)
                                continue

                            if not para.runs:
                                report["errors"].append({"key": cell_key, "error": "paragraph has no runs"})
                                continue

                            has_non_text = False
                            for run in para.runs:
                                if _has_non_text_content(run._element):
                                    has_non_text = True
                                    break

                            if has_non_text:
                                report["skipped_non_text_content"].append(cell_key)
                                continue

                            try:
                                if tc_writer:
                                    tc_writer.apply_paragraph_change(para, old_text, new_text)
                                else:
                                    if len(para.runs) > 1:
                                        report["format_simplified"].append(cell_key)
                                    _replace_paragraph_text(para, new_text)
                                report["applied"] += 1
                            except Exception as e:
                                logger.warning(f"Failed to apply modification to {cell_key}: {e}")
                                report["errors"].append({"key": cell_key, "error": str(e)})

        # Register rsid if Track Changes was used
        if tc_writer and report["applied"] > 0:
            tc_writer.register_rsid(doc)

        # Report unknown keys
        for key in mod_lookup:
            if key not in seen_keys:
                report["skipped_unknown_key"].append(key)

        report["success"] = report["applied"] > 0 or len(mod_lookup) == 0

        logger.info(
            f"Applied {report['applied']}/{len(mod_lookup)} modifications "
            f"(skipped: {len(report['skipped_not_editable'])} not-editable, "
            f"{len(report['skipped_unknown_key'])} unknown-key, "
            f"{len(report['skipped_no_change'])} no-change)"
        )

        output = BytesIO()
        doc.save(output)
        return output.getvalue(), report


def _replace_paragraph_text(para, new_text: str) -> None:
    """Replace paragraph text in-place by modifying w:t elements.

    Sets new_text on the first run, clears text on remaining runs.
    Preserves all XML structure (bookmarks, pPr, rPr, etc.).
    """
    runs = para.runs
    if not runs:
        return

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


# Singleton
_service: DocxRoundtripService | None = None


def get_docx_roundtrip_service() -> DocxRoundtripService:
    global _service
    if _service is None:
        _service = DocxRoundtripService()
    return _service
