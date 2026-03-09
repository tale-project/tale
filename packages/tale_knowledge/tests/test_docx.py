"""Tests for DOCX extraction."""

import pytest

from tale_knowledge.extraction.docx import extract_text_from_docx_bytes


def _make_simple_docx(text: str = "Hello World") -> bytes:
    """Create a minimal DOCX with text content."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractTextFromDocxBytes:
    @pytest.mark.asyncio
    async def test_basic_text_extraction(self):
        docx_bytes = _make_simple_docx("Hello World")
        text, vision_used, _ = await extract_text_from_docx_bytes(docx_bytes)
        assert "Hello World" in text
        assert vision_used is False

    @pytest.mark.asyncio
    async def test_multiple_paragraphs(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "Third paragraph" in text

    @pytest.mark.asyncio
    async def test_table_extraction(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        assert "A1" in text
        assert "B1" in text
        assert "[Table]" in text

    @pytest.mark.asyncio
    async def test_no_vision_without_client(self):
        docx_bytes = _make_simple_docx("No vision needed")
        text, vision_used, _ = await extract_text_from_docx_bytes(docx_bytes, vision_client=None)
        assert "No vision needed" in text
        assert vision_used is False

    @pytest.mark.asyncio
    async def test_header_extraction(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "My Header"
        doc.add_paragraph("Body text")
        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        assert "[Header]" in text
        assert "My Header" in text
        assert text.index("[Header]") < text.index("Body text")

    @pytest.mark.asyncio
    async def test_footer_extraction(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = "My Footer"
        doc.add_paragraph("Body text")
        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        assert "[Footer]" in text
        assert "My Footer" in text
        assert text.index("Body text") < text.index("[Footer]")

    @pytest.mark.asyncio
    async def test_header_and_footer_ordering(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = "Doc Header"
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = "Doc Footer"
        doc.add_paragraph("Middle content")
        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        header_pos = text.index("Doc Header")
        body_pos = text.index("Middle content")
        footer_pos = text.index("Doc Footer")
        assert header_pos < body_pos < footer_pos

    @pytest.mark.asyncio
    async def test_duplicate_headers_deduplicated(self):
        from io import BytesIO

        from docx import Document

        doc = Document()
        doc.add_paragraph("Page one")

        doc.add_section()
        section1 = doc.sections[0]
        section1.header.is_linked_to_previous = False
        section1.header.paragraphs[0].text = "Same Header"

        section2 = doc.sections[1]
        section2.header.is_linked_to_previous = False
        section2.header.paragraphs[0].text = "Same Header"

        buf = BytesIO()
        doc.save(buf)

        text, _, _ = await extract_text_from_docx_bytes(buf.getvalue())
        assert text.count("[Header]\nSame Header") == 1


class TestPageBreakDetection:
    @pytest.mark.asyncio
    async def test_no_page_breaks(self):
        docx_bytes = _make_simple_docx("Simple text")
        _, _, page_breaks = await extract_text_from_docx_bytes(docx_bytes)
        assert page_breaks == []

    @pytest.mark.asyncio
    async def test_inline_page_break(self):
        from io import BytesIO
        from lxml import etree

        from docx import Document

        doc = Document()
        doc.add_paragraph("Page one content")
        p = doc.add_paragraph()
        run = p.add_run("Page two content")
        br = etree.SubElement(
            run._element,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br",
        )
        br.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type",
            "page",
        )
        doc.add_paragraph("After break")
        buf = BytesIO()
        doc.save(buf)

        text, _, page_breaks = await extract_text_from_docx_bytes(buf.getvalue())
        assert len(page_breaks) >= 1
        assert "Page one content" in text
        assert "After break" in text

    @pytest.mark.asyncio
    async def test_page_break_before_style(self):
        from io import BytesIO
        from lxml import etree

        from docx import Document

        WP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        doc = Document()
        doc.add_paragraph("First page")
        p = doc.add_paragraph("Second page")
        pPr = p._element.find(f"{WP}pPr")
        if pPr is None:
            pPr = etree.SubElement(p._element, f"{WP}pPr")
            p._element.insert(0, pPr)
        etree.SubElement(pPr, f"{WP}pageBreakBefore")
        buf = BytesIO()
        doc.save(buf)

        text, _, page_breaks = await extract_text_from_docx_bytes(buf.getvalue())
        assert len(page_breaks) >= 1
        assert "First page" in text
        assert "Second page" in text

    @pytest.mark.asyncio
    async def test_multiple_page_breaks(self):
        from io import BytesIO
        from lxml import etree

        from docx import Document

        WP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        doc = Document()
        doc.add_paragraph("Page 1")

        p2 = doc.add_paragraph("Page 2")
        run2 = p2.runs[0]._element if p2.runs else p2.add_run("Page 2")._element
        br2 = etree.SubElement(run2, f"{WP}br")
        br2.set(f"{WP}type", "page")

        doc.add_paragraph("Still page 2")

        p3 = doc.add_paragraph("Page 3")
        run3 = p3.runs[0]._element if p3.runs else p3.add_run("Page 3")._element
        br3 = etree.SubElement(run3, f"{WP}br")
        br3.set(f"{WP}type", "page")

        buf = BytesIO()
        doc.save(buf)

        _, _, page_breaks = await extract_text_from_docx_bytes(buf.getvalue())
        assert len(page_breaks) == 2

    @pytest.mark.asyncio
    async def test_text_unchanged_with_page_breaks(self):
        from io import BytesIO
        from lxml import etree

        from docx import Document

        WP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        doc_no_breaks = Document()
        doc_no_breaks.add_paragraph("Hello")
        doc_no_breaks.add_paragraph("World")
        buf1 = BytesIO()
        doc_no_breaks.save(buf1)

        doc_with_breaks = Document()
        doc_with_breaks.add_paragraph("Hello")
        p = doc_with_breaks.add_paragraph("World")
        pPr = p._element.find(f"{WP}pPr")
        if pPr is None:
            pPr = etree.SubElement(p._element, f"{WP}pPr")
            p._element.insert(0, pPr)
        etree.SubElement(pPr, f"{WP}pageBreakBefore")
        buf2 = BytesIO()
        doc_with_breaks.save(buf2)

        text1, _, breaks1 = await extract_text_from_docx_bytes(buf1.getvalue())
        text2, _, breaks2 = await extract_text_from_docx_bytes(buf2.getvalue())

        assert text1 == text2
        assert breaks1 == []
        assert len(breaks2) >= 1
