"""Smart DOCX text extraction with Vision API support.

Extracts text and images from DOCX files while preserving the relative
positions of text and images in the document.
"""

from __future__ import annotations

import asyncio
import zipfile
from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from loguru import logger

from ._helpers import describe_image_bytes, extract_table_text

_WP_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

if TYPE_CHECKING:
    from tale_knowledge.vision.client import VisionClient

MAX_UNCOMPRESSED_SIZE = 500 * 1024 * 1024  # 500 MB


def _has_page_break(element) -> bool:
    """Detect explicit page breaks in a paragraph XML element.

    Checks for:
    - pageBreakBefore paragraph property
    - w:br with type="page" inline break
    """
    pPr = element.find(f"{_WP_NS}pPr")
    if pPr is not None:
        pb = pPr.find(f"{_WP_NS}pageBreakBefore")
        if pb is not None:
            val = pb.get(f"{_WP_NS}val")
            if val is None or val not in ("0", "false"):
                return True
    for br in element.iter(f"{_WP_NS}br"):
        if br.get(f"{_WP_NS}type") == "page":
            return True
    return False


async def extract_text_from_docx_bytes(
    docx_bytes: bytes,
    filename: str = "document.docx",
    *,
    vision_client: VisionClient | None = None,
    process_images: bool = True,
    max_concurrent: int = 3,
) -> tuple[str, bool, list[int]]:
    """Extract text from DOCX bytes with optional Vision support.

    Args:
        docx_bytes: Raw DOCX bytes.
        filename: Filename for logging.
        vision_client: Optional VisionClient for image description.
        process_images: Whether to extract and describe embedded images.
        max_concurrent: Max concurrent Vision API calls.

    Returns:
        Tuple of (extracted_text, vision_was_used, page_break_positions).
        page_break_positions is a list of element position indices where
        explicit page breaks occur. The text output is not modified.
    """
    logger.info(f"Processing DOCX: {filename}")

    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
            total = sum(info.file_size for info in zf.infolist())
            if total > MAX_UNCOMPRESSED_SIZE:
                raise ValueError(f"File exceeds maximum decompressed size ({total} bytes)")
    except zipfile.BadZipFile:
        raise ValueError("Invalid or corrupt file")

    doc = Document(BytesIO(docx_bytes))
    elements: list[tuple[int, str]] = []
    vision_used = False
    position = 0
    semaphore = asyncio.Semaphore(max_concurrent)

    image_rels: dict[str, bytes] = {}
    if process_images and vision_client:
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                try:
                    image_rels[rel.rId] = rel.target_part.blob
                except Exception as e:
                    logger.warning(f"Failed to extract image rel {rel.rId}: {e}")

    processed_image_rids: set[str] = set()
    page_break_positions: list[int] = []

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            if _has_page_break(element):
                page_break_positions.append(position)

            para = para_map.get(element)

            if para:
                text = para.text.strip()
                if text:
                    elements.append((position, text))
                    position += 1

                if process_images and vision_client:
                    blip_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
                    embed_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

                    for blip in element.iter(f"{blip_ns}blip"):
                        embed_id = blip.get(f"{embed_ns}embed")
                        if embed_id and embed_id in image_rels and embed_id not in processed_image_rids:
                            processed_image_rids.add(embed_id)
                            image_bytes = image_rels[embed_id]
                            description = await describe_image_bytes(image_bytes, semaphore, vision_client)
                            if description:
                                elements.append((position, f"[Image: {description}]"))
                                vision_used = True
                                position += 1

        elif tag == "tbl":
            table = table_map.get(element)
            if table:
                table_text = extract_table_text(table)
                if table_text:
                    elements.append((position, f"[Table]\n{table_text}"))
                    position += 1

    if process_images and vision_client:
        for rid, image_bytes in image_rels.items():
            if rid not in processed_image_rids:
                description = await describe_image_bytes(image_bytes, semaphore, vision_client)
                if description:
                    elements.append((position, f"[Image: {description}]"))
                    vision_used = True
                    position += 1

    header_texts: list[str] = []
    footer_texts: list[str] = []
    seen_headers: set[str] = set()
    seen_footers: set[str] = set()

    for section in doc.sections:
        for hdr in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if hdr and not hdr.is_linked_to_previous:
                text = "\n".join(p.text.strip() for p in hdr.paragraphs if p.text.strip())
                if text and text not in seen_headers:
                    seen_headers.add(text)
                    header_texts.append(text)

        for ftr in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if ftr and not ftr.is_linked_to_previous:
                text = "\n".join(p.text.strip() for p in ftr.paragraphs if p.text.strip())
                if text and text not in seen_footers:
                    seen_footers.add(text)
                    footer_texts.append(text)

    header_position = -len(header_texts)
    for hdr_text in header_texts:
        elements.append((header_position, f"[Header]\n{hdr_text}"))
        header_position += 1

    for ftr_text in footer_texts:
        elements.append((position, f"[Footer]\n{ftr_text}"))
        position += 1

    elements.sort(key=lambda x: x[0])
    content = "\n\n".join(elem[1] for elem in elements)

    logger.info(f"DOCX processing complete: {len(elements)} elements, Vision API used: {vision_used}")

    return content, vision_used, page_break_positions
